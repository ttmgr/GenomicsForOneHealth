"""Export all Mann-Whitney U tests to a DOCX report.

Source of truth:
- outputs/optional/stats_all.csv

This report includes only rows where stat_name == "U" and excludes
Spearman and Kruskal-Wallis results.
"""

from __future__ import annotations

import argparse
from collections import OrderedDict
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from listeria_pipeline_common import (
    DEFAULT_INPUT_PATH,
    DEFAULT_INPUT_SHEET,
    DEFAULT_OUTPUT_ROOT,
    ROOT,
    build_pipeline_paths,
    format_input_label,
)

INPUT_STATS = DEFAULT_OUTPUT_ROOT / "optional" / "stats_all.csv"
OUTPUT_DOCX = DEFAULT_OUTPUT_ROOT / "publication_v4" / "listeria_mann_whitney_u_methods_and_results.docx"
SOURCE_INPUT = DEFAULT_INPUT_PATH
SOURCE_SHEET = DEFAULT_INPUT_SHEET

FIGURE_DESCRIPTIONS = OrderedDict(
    [
        (
            "Fig 3a.3 (AS vs N)",
            "Sponge quasimetagenomic samples, stratified by timepoint and inoculum level "
            "(Lm2, Lm4, Lm6), comparing Native (N) versus Adaptive Sampling (AS).",
        ),
        (
            "Fig 3b.2 (AS vs N)",
            "Metagenomic samples, stratified by swab type, comparing Native (N) versus "
            "Adaptive Sampling (AS).",
        ),
        (
            "Fig 3b.1 (swab pairwise)",
            "Metagenomic N-only samples, pairwise swab-type comparisons.",
        ),
        (
            "Fig 3a.5 (dose)",
            "Sponge quasimetagenomic N-only samples at pooled 24 h, pairwise inoculum "
            "comparisons among Lm2, Lm4, and Lm6.",
        ),
        (
            "Fig 3c (kit)",
            "Quasimetagenomic N-only samples, Micro versus Mini extraction kit "
            "comparisons pooled across swab types and within each swab type.",
        ),
        (
            "Fig 3a.3p (pooled AS vs N)",
            "Sponge quasimetagenomic samples with Lm2/Lm4/Lm6 pooled, comparing Native "
            "(N) versus Adaptive Sampling (AS) by timepoint.",
        ),
        (
            "Fig 3d (baseline vs time)",
            "Sponge N-only pooled analyses comparing metagenomic baseline (0 h) against "
            "later quasimetagenomic timepoints, including a pooled 24 h comparison.",
        ),
        (
            "Fig 3b.1p (pooled quasi swab)",
            "Quasimetagenomic N-only samples with Lm2/Lm4/Lm6 pooled, pairwise swab-type "
            "comparisons.",
        ),
    ]
)

METRIC_ORDER = {
    "Lm % Reads": 0,
    "Lm % Bases": 1,
    "L. ivanovii Nr26 Proportion": 2,
}


def configure_paths(
    output_root: str | Path | None = None,
    source_input: str | Path | None = None,
    source_sheet: str = DEFAULT_INPUT_SHEET,
) -> None:
    global INPUT_STATS, OUTPUT_DOCX, SOURCE_INPUT, SOURCE_SHEET

    paths = build_pipeline_paths(output_root=output_root)
    INPUT_STATS = paths.optional_dir / "stats_all.csv"
    OUTPUT_DOCX = paths.publication_dir / "listeria_mann_whitney_u_methods_and_results.docx"
    SOURCE_INPUT = Path(source_input or DEFAULT_INPUT_PATH)
    SOURCE_SHEET = source_sheet


def display_path(path: Path) -> str:
    candidate = path.resolve() if not path.is_absolute() else path
    try:
        return str(candidate.relative_to(ROOT))
    except ValueError:
        return str(candidate)


def fmt_num(value: float | int | str) -> str:
    if pd.isna(value):
        return "n/a"
    if isinstance(value, str):
        return value
    return f"{float(value):.4g}"


def fmt_p(value: float | int | str) -> str:
    if pd.isna(value):
        return "n/a"
    if isinstance(value, str):
        return value
    return f"{float(value):.3g}"


def set_run_font(run, size: float = 9, bold: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold


def add_text(paragraph, text, size: float = 10, bold: bool = False) -> None:
    if text is None or (isinstance(text, float) and pd.isna(text)):
        text = ""
    elif not isinstance(text, str):
        text = str(text)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    add_text(p, text, size=10)


def derive_groups(row: pd.Series) -> tuple[str, str]:
    comparison = str(row["comparison"])
    figure = str(row["figure"])

    if comparison.startswith("Micro vs Mini"):
        return "Micro", "Mini"
    if comparison.startswith("0h vs "):
        return "0h", comparison.replace("0h vs ", "", 1).strip()
    if comparison.endswith(" (N only)"):
        left, right = comparison.replace(" (N only)", "").split(" vs ", 1)
        return left.strip(), right.strip()
    if " vs " in comparison:
        left, right = comparison.split(" vs ", 1)
        return left.strip(), right.strip()
    return "Group 1", "Group 2"


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def build_summary_table(document: Document, df: pd.DataFrame) -> None:
    summary = (
        df.groupby("figure", sort=False)
        .agg(
            tests=("figure", "size"),
            raw_sig=("p", lambda s: int((s < 0.05).sum())),
            bh_sig=("p_BH", lambda s: int((s < 0.05).sum())),
        )
        .reset_index()
    )

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    headers = ["Figure", "MWU tests", "Raw p < 0.05", "BH-adjusted p < 0.05"]
    for cell, text in zip(header, headers):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, text, size=9, bold=True)

    for _, row in summary.iterrows():
        cells = table.add_row().cells
        values = [row["figure"], str(row["tests"]), str(row["raw_sig"]), str(row["bh_sig"])]
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, value, size=8.5)


def build_results_table(document: Document, fig_name: str, df: pd.DataFrame) -> None:
    cols = [
        "Subgroup",
        "Metric",
        "Group 1",
        "Group 2",
        "n1",
        "n2",
        "Median 1",
        "Median 2",
        "U",
        "p",
        "p_BH",
        "sig",
        "sig_BH",
    ]
    table = document.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0].cells
    for cell, text in zip(header, cols):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, text, size=8.5, bold=True)

    sub = df[df["figure"] == fig_name].copy()
    sub["metric_order"] = sub["metric"].map(METRIC_ORDER).fillna(99)
    sub = sub.sort_values(["metric_order", "comparison", "subgroup"]).drop(columns=["metric_order"])

    for _, row in sub.iterrows():
        cells = table.add_row().cells
        values = [
            row["subgroup"],
            row["metric"],
            row["group_1"],
            row["group_2"],
            str(int(row["n1"])),
            str(int(row["n2"])),
            fmt_num(row["median1"]),
            fmt_num(row["median2"]),
            fmt_num(row["stat"]),
            fmt_p(row["p"]),
            fmt_p(row["p_BH"]),
            row["sig"],
            row["sig_BH"],
        ]
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, value, size=8)


def main() -> None:
    args = parse_args()
    configure_paths(
        output_root=args.output_root,
        source_input=args.source_input,
        source_sheet=args.source_sheet,
    )
    df = pd.read_csv(INPUT_STATS)
    df = df[df["stat_name"] == "U"].copy()
    if df.empty:
        raise ValueError("No Mann-Whitney U rows found in stats_all.csv")

    groups = df.apply(derive_groups, axis=1, result_type="expand")
    df["group_1"] = groups[0]
    df["group_2"] = groups[1]

    document = Document()
    set_landscape(document.sections[0])

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)

    title = document.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(title, "Listeria Mann-Whitney U Tests: Methods and Results", size=16, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(
        subtitle,
        f"Generated {date.today().isoformat()} from {display_path(INPUT_STATS)}",
        size=9,
    )

    document.add_heading("Scope", level=1)
    p = document.add_paragraph()
    add_text(
        p,
        "This document contains every Mann-Whitney U result stored in the combined "
        "statistics table. It excludes Spearman correlation and Kruskal-Wallis rows "
        "because those are not Mann-Whitney U tests.",
        size=10,
    )

    total_tests = len(df)
    raw_sig = int((df["p"] < 0.05).sum())
    bh_sig = int((df["p_BH"] < 0.05).sum())
    add_bullet(document, f"Total Mann-Whitney U tests: {total_tests}")
    add_bullet(document, f"Raw p < 0.05: {raw_sig}")
    add_bullet(document, f"BH-adjusted p < 0.05: {bh_sig}")

    document.add_heading("Methods", level=1)
    add_bullet(
        document,
        f"Primary data source: {format_input_label(SOURCE_INPUT, SOURCE_SHEET)}, processed by "
        "scripts/analyze_listeria.py and exported to the run-specific optional/stats_all.csv.",
    )
    add_bullet(
        document,
        "Data cleaning in the source pipeline: headers were stripped, timepoints were "
        "mapped from sample-type labels, an included_in_analysis flag was added, and "
        "rows with Treatment = Blank or Background were retained in the master workbook "
        "but excluded from all downstream tests and plots.",
    )
    add_bullet(
        document,
        "Metagenomic rows retained source Treatment labels in the master data when present, "
        "but the main time-course analyses still pooled metagenomic rows as the shared 0h baseline.",
    )
    add_bullet(
        document,
        "Mann-Whitney U tests were run with scipy.stats.mannwhitneyu(a, b, "
        "alternative='two-sided') after dropping missing values within each comparison.",
    )
    add_bullet(
        document,
        "This appendix reports U, raw p, BH-adjusted p (p_BH), sample sizes (n1, n2), "
        "and group medians (Median 1, Median 2) for every retained MWU row.",
    )
    add_bullet(
        document,
        "The BH-adjusted p-values were inherited from the original analysis blocks in "
        "scripts/analyze_listeria.py before those blocks were concatenated into the "
        "combined CSV. They are therefore family-specific adjusted p-values, not a fresh "
        "single correction across all 111 MWU rows in this document.",
    )
    add_bullet(
        document,
        "For figure families labeled 'AS vs N', Group 1 corresponds to N and Group 2 "
        "corresponds to AS, because the underlying code passed N first and AS second "
        "into the test function even though the human-readable label says 'AS vs N'.",
    )
    add_bullet(
        document,
        "For other contrasts, Group 1 and Group 2 follow the left-to-right comparison "
        "definition (for example, Sponge vs Cotton, Micro vs Mini, or 0h vs 24h pooled).",
    )
    add_bullet(
        document,
        "Significance coding: *** p < 0.001, ** p < 0.01, * p < 0.05, ns otherwise. "
        "The 'sig' column uses raw p, and 'sig_BH' uses BH-adjusted p_BH.",
    )

    document.add_heading("Counts by Figure", level=1)
    build_summary_table(document, df)

    document.add_heading("Column Guide", level=1)
    add_bullet(document, "Subgroup: the figure-specific analysis stratum for that test.")
    add_bullet(document, "Metric: response variable tested in that row.")
    add_bullet(document, "Group 1 / Group 2: actual group identities passed into the MWU test.")
    add_bullet(document, "n1 / n2: non-missing sample counts for Group 1 and Group 2.")
    add_bullet(document, "Median 1 / Median 2: medians for Group 1 and Group 2.")
    add_bullet(document, "U: Mann-Whitney U statistic.")
    add_bullet(document, "p: raw two-sided p-value.")
    add_bullet(document, "p_BH: stored Benjamini-Hochberg adjusted p-value from the original analysis family.")

    document.add_heading("Results Tables", level=1)
    for fig_name, description in FIGURE_DESCRIPTIONS.items():
        if fig_name not in set(df["figure"]):
            continue
        document.add_heading(fig_name, level=2)
        p = document.add_paragraph()
        add_text(p, description, size=10)
        build_results_table(document, fig_name, df)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    print(f"Saved {OUTPUT_DOCX}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Export the Mann-Whitney U appendix DOCX for a Listeria analysis run."
    )
    parser.add_argument(
        "--output-root",
        type=Path,
        default=DEFAULT_OUTPUT_ROOT,
        help="Root output directory containing optional/stats_all.csv and publication_v4/.",
    )
    parser.add_argument(
        "--source-input",
        type=Path,
        default=DEFAULT_INPUT_PATH,
        help="Original input data file used for the analysis run.",
    )
    parser.add_argument(
        "--source-sheet",
        default=DEFAULT_INPUT_SHEET,
        help="Worksheet name for XLSX source inputs. Ignored for CSV.",
    )
    return parser.parse_args()


if __name__ == "__main__":
    main()
